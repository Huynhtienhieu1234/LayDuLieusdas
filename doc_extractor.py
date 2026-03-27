"""
doc_extractor.py
----------------
Đọc file .docx (đề cương hoặc bài giảng) và tự động trích xuất:
  - Tổng quan về nội dung môn học
  - Mục tiêu học tập
  - Yêu cầu đối với học viên
  - Phương pháp giảng dạy và đánh giá

Không cần AI, không cần internet.

Cài đặt:
    pip install python-docx

Chạy:
    python doc_extractor.py <file.docx>
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("❌  Thiếu thư viện: pip install python-docx")
    sys.exit(1)


# ══════════════════════════════════════════════════════════════
# KEYWORD PATTERNS — thêm từ khóa nếu file của bạn dùng tên khác
# ══════════════════════════════════════════════════════════════

SECTION_PATTERNS = {
    "tong_quan": [
        r"lời giới thiệu",
        r"^lời nói đầu",
        r"thông tin chung",
        r"tổng quan",
        r"mô tả môn học",
        r"giới thiệu môn",
    ],
    "muc_tieu": [
        r"mục tiêu",
        r"mục đích (học phần|môn học|bài giảng)",
    ],
    "yeu_cau": [
        r"yêu cầu (chung )?đối với (sinh|học) viên",
        r"yêu cầu chung",
        r"điều kiện (tham gia|tiên quyết)",
    ],
    "phuong_phap": [
        r"kiểm tra và đánh giá",
        r"phương pháp đánh giá",
        r"hình thức (kiểm tra|đánh giá)",
        r"^5[\.\s].*đánh giá",
    ],
}

# Các heading dùng để nhận biết "phần mới bắt đầu → dừng lấy nội dung"
STOP_PATTERNS = [
    r"^\d+[\.\d]*[\s\.]",          # "1.", "2.1.", "5. "
    r"^chương \d",                  # "Chương 1"
    r"^(lời giới thiệu|mục lục|danh mục|phụ lục|tài liệu)",
]


def read_docx(path: str) -> list[str]:
    """Đọc toàn bộ đoạn văn từ paragraphs + tables, theo thứ tự xuất hiện."""
    doc = Document(path)
    seen = set()
    lines = []

    def add(text):
        t = text.strip()
        if t and t not in seen:
            seen.add(t)
            lines.append(t)

    for p in doc.paragraphs:
        add(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    add(p.text)

    return lines


def norm(text: str) -> str:
    return text.lower().strip()


def matches(text: str, patterns: list[str]) -> bool:
    n = norm(text)
    return any(re.search(pat, n) for pat in patterns)


def is_stop_line(text: str) -> bool:
    n = norm(text)
    return any(re.search(pat, n) for pat in STOP_PATTERNS)


def extract_after(lines: list[str], patterns: list[str], max_lines: int = 8) -> list[str]:
    """
    Tìm dòng tiêu đề khớp patterns, thu thập các dòng nội dung ngay sau.
    Dừng khi gặp tiêu đề mới hoặc đủ max_lines.
    """
    result = []
    capturing = False
    count = 0
    all_patterns = [p for ps in SECTION_PATTERNS.values() for p in ps]

    for line in lines:
        if matches(line, patterns):
            capturing = True
            count = 0
            # Nếu tiêu đề chứa nội dung sau dấu ":" thì lấy luôn
            if ":" in line:
                after = line.split(":", 1)[1].strip()
                if len(after) > 15:
                    result.append(after)
                    count += 1
            continue

        if capturing:
            if count >= max_lines:
                break
            if is_stop_line(line) or matches(line, all_patterns):
                break
            if len(line) < 6:
                continue
            result.append(line)
            count += 1

    return result


# ══════════════════════════════════════════════════════════════
# TRÍCH XUẤT TỪNG PHẦN
# ══════════════════════════════════════════════════════════════

def get_tong_quan(lines: list[str]) -> str:
    # Ưu tiên 1: lấy đoạn lời giới thiệu
    result = extract_after(lines, SECTION_PATTERNS["tong_quan"], max_lines=4)
    if result:
        return " ".join(result)

    # Ưu tiên 2: ghép thông tin cơ bản (tên bài giảng, đối tượng, tín chỉ)
    info = []
    for line in lines[:40]:
        n = norm(line)
        if re.search(r"tên (bài giảng|học phần|môn)", n):
            info.append(line)
        elif re.search(r"(đối tượng|số tín chỉ|thời lượng)", n):
            info.append(line)
    if info:
        return " | ".join(info[:5])

    return "[Không trích xuất được tổng quan]"


def get_muc_tieu(lines: list[str]) -> list[str]:
    result = extract_after(lines, SECTION_PATTERNS["muc_tieu"], max_lines=8)

    # Lọc bỏ sub-heading "2.1. Mục tiêu chung", "2.2. Mục tiêu cụ thể"
    cleaned = []
    for line in result:
        n = norm(line)
        if re.match(r"^\d+\.\d+", line) and "mục tiêu" in n:
            continue
        if re.search(r"^(mục tiêu chung|mục tiêu cụ thể)$", n):
            continue
        cleaned.append(line)

    return cleaned or ["[Không trích xuất được mục tiêu]"]


def get_yeu_cau(lines: list[str]) -> list[str]:
    result = extract_after(lines, SECTION_PATTERNS["yeu_cau"], max_lines=6)

    # Nếu không lấy được, tìm bullet "-" trong bảng cấu trúc
    if not result:
        for i, line in enumerate(lines):
            if re.search(r"yêu cầu", norm(line)):
                for l in lines[i+1:i+8]:
                    stripped = l.lstrip("-•+ \t")
                    if len(stripped) > 10:
                        result.append(stripped)
                if result:
                    break

    # Chuẩn hóa: bỏ dấu "-" đầu dòng
    cleaned = [re.sub(r"^[-•+]\s*", "", l).strip() for l in result]
    return [l for l in cleaned if l] or ["[Không trích xuất được yêu cầu]"]


def get_phuong_phap(lines: list[str]) -> list[str]:
    result = extract_after(lines, SECTION_PATTERNS["phuong_phap"], max_lines=7)

    # Bỏ sub-heading "5.1.", "5.2."
    cleaned = []
    for line in result:
        if re.match(r"^\d+\.\d+", line) and len(line) < 40:
            continue
        cleaned.append(line)

    return cleaned or ["[Không trích xuất được phương pháp đánh giá]"]


# ══════════════════════════════════════════════════════════════
# IN KẾT QUẢ
# ══════════════════════════════════════════════════════════════

def print_result(tong_quan, muc_tieu, yeu_cau, phuong_phap):
    SEP = "─" * 62

    print(f"\n{SEP}")
    print("Tổng quan về nội dung môn học:")
    print(f"  {tong_quan}")

    print(f"\n{SEP}")
    print("Mục tiêu học tập:")
    for item in muc_tieu:
        print(f"  • {item}")

    print(f"\n{SEP}")
    print("Yêu cầu đối với học viên:")
    for item in yeu_cau:
        print(f"  • {item}")

    print(f"\n{SEP}")
    print("Phương pháp giảng dạy và đánh giá:")
    for item in phuong_phap:
        print(f"  • {item}")

    print(f"{SEP}\n")


# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════

def main():
    if len(sys.argv) < 2:
        print("Cách dùng: python doc_extractor.py <file.docx>")
        print("Ví dụ:     python doc_extractor.py De_cuong.docx")
        sys.exit(1)

    path = sys.argv[1]
    if not Path(path).exists():
        print(f"❌  Không tìm thấy file: {path}")
        sys.exit(1)

    print(f"\n📖  Đang đọc: {path}")
    lines = read_docx(path)
    print(f"✅  Đọc xong — {len(lines)} đoạn văn bản")

    tong_quan   = get_tong_quan(lines)
    muc_tieu    = get_muc_tieu(lines)
    yeu_cau     = get_yeu_cau(lines)
    phuong_phap = get_phuong_phap(lines)

    print_result(tong_quan, muc_tieu, yeu_cau, phuong_phap)


if __name__ == "__main__":
    main()
